from docx import Document
import json

# Read the resume
doc = Document('resumeDeepak.docx')

# Extract all text
all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text)

# Also extract from tables if any
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_text.append(cell.text)

# Print all content
for line in all_text:
    print(line)
